import os
from docx import Document

class DocumentReader:
  def __init__(self, file_name):
    self.document = self.read_document(file_name)

  def read_document(self, file_name):
    current_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(current_dir, file_name)
    return Document(file_path)
  
  def get_hazards(self):
    hazards = {}
    current_category = None
    current_sub_category = None
    for table in self.document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    style = paragraph.style.name

                    # Detect main category
                    if style == "Heading 1":
                        current_category = paragraph.text.strip()
                        hazards[current_category] = {}

                    # Detect sub-category
                    elif style == "Heading 2" and current_category:
                        current_sub_category = paragraph.text.strip()
                        hazards[current_category][current_sub_category] = {"To Whom": [], "Controls": []}

                    # Detect "Normal" paragraphs
                    elif style == "normal" and current_sub_category:
                        is_bold = paragraph.runs[0].bold
                        text = paragraph.text.strip()

                        # if text is bold it is a "To Whom" paragraph, else it is control
                        if text != "" and is_bold:
                            hazards[current_category][current_sub_category]["To Whom"].append(text)
                        elif text !="" and not is_bold:
                            hazards[current_category][current_sub_category]["Controls"].append(text)

    return hazards

