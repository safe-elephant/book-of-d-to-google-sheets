import os

from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


def read_document(file_name):
    current_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(current_dir, file_name)
    return Document(file_path)

def get_top_level_categories(document):
    unique_top_level_categories = []
    seen = set()

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.style.name == "Heading 1":
                        text = paragraph.text.strip()
                        if text not in seen:
                            unique_top_level_categories.append(text)
                            seen.add(text)
                        break  # Assuming only one 'Heading 1' per cell

    return unique_top_level_categories

document = read_document("./book-of-d.docx")
top_level_categories = get_top_level_categories(document)

x = top_level_categories
